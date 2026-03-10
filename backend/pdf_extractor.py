import PyPDF2
from io import BytesIO
from typing import Union, Dict, Any
import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_file: Union[BytesIO, bytes]) -> str:
    """
    Extract text content from a PDF file with comprehensive error handling.
    
    Args:
        pdf_file: PDF file object (from Streamlit file uploader) or bytes
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If PDF is invalid or corrupted
        TypeError: If input type is invalid
    """
    if not pdf_file:
        raise ValueError("PDF file cannot be None or empty")
    
    try:
        # Create PDF reader object
        if isinstance(pdf_file, bytes):
            pdf_file = BytesIO(pdf_file)
            
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF has no pages")
        
        # Extract text from all pages
        text_parts = []
        for i, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {i+1}: {e}")
                continue
        
        text = "\n".join(text_parts).strip()
        
        if not text:
            raise ValueError("Could not extract any text from PDF. The file might be image-based or corrupted.")
        
        logger.info(f"Successfully extracted {len(text)} characters from PDF with {len(pdf_reader.pages)} pages")
        return text
        
    except PyPDF2.errors.PdfReadError as e:
        raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error extracting PDF: {e}")
        raise ValueError(f"Error extracting text from PDF: {str(e)}")


def get_pdf_metadata(pdf_file: Union[BytesIO, bytes]) -> Dict[str, Any]:
    """
    Extract metadata from a PDF file with error handling.
    
    Args:
        pdf_file: PDF file object
        
    Returns:
        Dictionary containing PDF metadata
        
    Raises:
        ValueError: If PDF is invalid
    """
    if not pdf_file:
        raise ValueError("PDF file cannot be None or empty")
    
    try:
        if isinstance(pdf_file, bytes):
            pdf_file = BytesIO(pdf_file)
            
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        metadata = {
            'num_pages': len(pdf_reader.pages),
            'author': 'Unknown',
            'title': 'Unknown',
            'subject': 'Unknown'
        }
        
        if pdf_reader.metadata:
            metadata.update({
                'author': pdf_reader.metadata.get('/Author', 'Unknown'),
                'title': pdf_reader.metadata.get('/Title', 'Unknown'),
                'subject': pdf_reader.metadata.get('/Subject', 'Unknown')
            })
        
        logger.info(f"Extracted metadata for PDF with {metadata['num_pages']} pages")
        return metadata
        
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {e}")
        return {'error': str(e), 'num_pages': 0, 'author': 'Unknown', 'title': 'Unknown', 'subject': 'Unknown'}


def extract_text_from_docx(docx_file: Union[BytesIO, bytes]) -> str:
    """
    Extract text content from a DOCX file with comprehensive error handling.

    Args:
        docx_file: DOCX file object (from Streamlit file uploader) or bytes

    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If DOCX is invalid or corrupted
        TypeError: If input type is invalid
    """
    if not docx_file:
        raise ValueError("DOCX file cannot be None or empty")
    
    try:
        if isinstance(docx_file, bytes):
            docx_file = BytesIO(docx_file)
            
        document = docx.Document(docx_file)
        paragraphs = []
        
        for para in document.paragraphs:
            if para.text and para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        text = "\n".join(paragraphs).strip()

        if not text:
            raise ValueError("Could not extract any text from DOCX. The file might be empty or corrupted.")

        logger.info(f"Successfully extracted {len(text)} characters from DOCX with {len(document.paragraphs)} paragraphs")
        return text
        
    except docx.opc.exceptions.PackageNotFoundError:
        raise ValueError("Invalid DOCX file format")
    except Exception as e:
        logger.error(f"Unexpected error extracting DOCX: {e}")
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")